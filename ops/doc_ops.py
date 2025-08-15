from pathlib import Path
from typing import Optional
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from pdfminer.high_level import extract_text

def pdf_to_docx(src: Path, out_docx: Path, ocr: bool = False, dpi: int = 200) -> Path:
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if ocr:
        # OCR: rasteriza páginas e reconhece com Tesseract
        images = convert_from_path(str(src), dpi=dpi)
        for img in images:
            text = pytesseract.image_to_string(img, lang="por+eng")
            if text.strip():
                for line in text.splitlines():
                    doc.add_paragraph(line)
            doc.add_page_break()
    else:
        # Extração de texto por PDFMiner (sem OCR)
        text = extract_text(str(src)) or ""
        for line in text.splitlines():
            doc.add_paragraph(line)
    doc.save(str(out_docx))
    return out_docx
