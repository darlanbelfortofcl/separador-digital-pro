import os, platform, subprocess
from pdf2docx import Converter
from docx import Document
from pdfminer.high_level import extract_text
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image
import pandas as pd
import pytesseract
from pdf2image import convert_from_path

def convert_pdf_to_docx(src_path, dst_path):
    cv = Converter(src_path)
    try:
        cv.convert(dst_path, start=0, end=None)
    finally:
        cv.close()

def convert_docx_to_pdf(src_path, dst_path):
    if platform.system().lower().startswith("win"):
        from docx2pdf import convert
        convert(src_path, dst_path)
        return
    if _has_libreoffice():
        _convert_with_libreoffice(src_path, dst_path, "pdf")
        return
    raise RuntimeError("DOCX→PDF requer Windows (docx2pdf) ou LibreOffice no servidor.")

def convert_docx_to_txt(src_path, dst_path):
    doc = Document(src_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    with open(dst_path, "w", encoding="utf-8") as f:
        f.write(text)

def convert_txt_to_docx(src_path, dst_path):
    with open(src_path, "r", encoding="utf-8") as f:
        text = f.read()
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(dst_path)

def convert_pdf_to_txt(src_path, dst_path):
    text = extract_text(src_path)
    with open(dst_path, "w", encoding="utf-8") as f:
        f.write(text)

def convert_txt_to_pdf(src_path, dst_path):
    with open(src_path, "r", encoding="utf-8") as f:
        text = f.read()
    c = canvas.Canvas(dst_path, pagesize=letter)
    width, height = letter
    x, y = 30, height - 40
    for line in text.split("\n"):
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(x, y, line[:1000])
        y -= 16
    c.save()

def convert_xlsx_to_csv(src_path, dst_path):
    df = pd.read_excel(src_path)
    df.to_csv(dst_path, index=False)

def convert_csv_to_xlsx(src_path, dst_path):
    df = pd.read_csv(src_path)
    df.to_excel(dst_path, index=False)

def convert_image_to_pdf(src_path, dst_path):
    image = Image.open(src_path)
    rgb = image.convert("RGB")
    rgb.save(dst_path)

def convert_scanned_pdf_to_docx(src_path, dst_path):
    images = convert_from_path(src_path)
    doc = Document()
    for img in images:
        text = pytesseract.image_to_string(img)
        doc.add_paragraph(text)
    doc.save(dst_path)

def _has_libreoffice():
    for cmd in ("soffice", "libreoffice"):
        try:
            subprocess.run([cmd, "--version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=False)
            return True
        except FileNotFoundError:
            continue
    return False

def _convert_with_libreoffice(src_path, dst_path, out_format):
    out_dir = os.path.dirname(dst_path)
    os.makedirs(out_dir, exist_ok=True)
    subprocess.run(["soffice", "--headless", "--convert-to", out_format, "--outdir", out_dir, src_path], check=True)

CONVERSION_FUNCTIONS = {
    "pdf2docx": convert_pdf_to_docx,
    "docx2pdf": convert_docx_to_pdf,
    "docx2txt": convert_docx_to_txt,
    "txt2docx": convert_txt_to_docx,
    "pdf2txt": convert_pdf_to_txt,
    "txt2pdf": convert_txt_to_pdf,
    "xlsx2csv": convert_xlsx_to_csv,
    "csv2xlsx": convert_csv_to_xlsx,
    "jpg2pdf": convert_image_to_pdf,
    "jpeg2pdf": convert_image_to_pdf,
    "png2pdf": convert_image_to_pdf,
    "pdf2docx_ocr": convert_scanned_pdf_to_docx
}
