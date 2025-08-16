from pathlib import Path
from typing import Iterable, List, Tuple
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTTextLine, LTAnno
from docx import Document
from docx.shared import Inches

def _iter_text_boxes(pdf_path: Path) -> Iterable[LTTextContainer]:
    for page_layout in extract_pages(str(pdf_path)):
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                yield element

def _line_text_and_bold(line: LTTextLine) -> tuple[str, bool]:
    total_chars = 0
    bold_chars = 0
    text_parts: List[str] = []
    for obj in line:
        if isinstance(obj, LTChar):
            total_chars += 1
            name = (obj.fontname or "").lower()
            if "bold" in name or "black" in name or "heavy" in name:
                bold_chars += 1
            text_parts.append(obj.get_text())
        elif isinstance(obj, LTAnno):
            text_parts.append(obj.get_text())
        else:
            try:
                text_parts.append(obj.get_text())
            except Exception:
                pass
    text = "".join(text_parts)
    is_bold = (bold_chars / total_chars) >= 0.5 if total_chars else False
    return text, is_bold

def pdf_to_docx_editable(pdf_path: Path, out_docx: Path) -> Path:
    doc = Document()
    doc.core_properties.title = pdf_path.stem
    current_para = None
    last_y = None

    for tb in _iter_text_boxes(pdf_path):
        for line in tb:
            if not isinstance(line, LTTextLine):
                continue
            text, is_bold = _line_text_and_bold(line)
            text = text.replace('\u00a0', ' ').strip()
            if not text:
                continue
            y0 = getattr(line, 'y0', None)
            if last_y is not None and y0 is not None and abs(y0 - last_y) > 12:
                current_para = None
            last_y = y0
            if current_para is None:
                current_para = doc.add_paragraph()
            run = current_para.add_run(text)
            if is_bold:
                run.bold = True
        current_para = None

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    return out_docx
